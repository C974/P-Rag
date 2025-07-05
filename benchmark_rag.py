## genereting answers and calculating similarties 
# #!/usr/bin/env python3
# """
# Palestine RAG Benchmarking Script

# This script evaluates the Palestine RAG system against a benchmark dataset containing
# questions with multiple choice answers and Bloom taxonomy classifications.

# Features:
# - Loads benchmark from JSONL file
# - Uses existing RAG retrieval and generation
# - Evaluates accuracy by comparing generated answers to choices
# - Provides detailed analysis by Bloom taxonomy level
# - Outputs comprehensive JSON and text reports
# """

# import json
# import jsonlines
# import os
# import sys
# import time
# from datetime import datetime
# from typing import List, Dict, Any, Tuple
# import re
# from collections import defaultdict, Counter
# import torch
# from sentence_transformers import SentenceTransformer
# import numpy as np
# import pickle
# from tqdm import tqdm
# import PyPDF2
# from docx import Document as DocxDocument
# import csv

# # Hugging Face imports
# from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline



# # Import from your existing P-RAG system
# from bloom_utils import bloom_levels, get_bloom_instruction

# # Configuration - matching your P-RAG.py
# SOURCES_FOLDER = "./Sources"
# LANGUAGE_MODEL = 'Qwen/Qwen3-1.7B'  # Using QCRI Fanar model from HuggingFace
# EMBEDDING_MODEL_NAME = 'intfloat/e5-large-v2'  # Better embedding model
# VECTOR_DB_FILE = "vector_db_cache_bge.pkl"  # Separate cache for BGE model

# # Benchmark configuration
# BENCHMARK_FILE = "converted_benchmark_with_bloom.jsonl"
# OUTPUT_DIR = "benchmark_results"
# MAX_RESPONSE_LENGTH = 1000
# RETRIEVAL_TOP_K = 5

# class RAGBenchmarker:
#     def __init__(self):
#         """Initialize the RAG benchmarker with existing embeddings and models"""
#         self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
#         print(f"Using device: {self.device}")
        
#         # Initialize embedding model
#         self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME, device=self.device)
        
#         # Initialize Hugging Face language model
#         print(f"Loading language model: {LANGUAGE_MODEL}")
#         try:
#             self.tokenizer = AutoTokenizer.from_pretrained(LANGUAGE_MODEL, trust_remote_code=True)
#             self.language_model = AutoModelForCausalLM.from_pretrained(
#                 LANGUAGE_MODEL,
#                 torch_dtype=torch.float16 if self.device == 'cuda' else torch.float32,
#                 device_map='auto' if self.device == 'cuda' else None,
#                 trust_remote_code=True
#             )
            
#             # Set padding token if not exists
#             if self.tokenizer.pad_token is None:
#                 self.tokenizer.pad_token = self.tokenizer.eos_token
            
#             # Create text generation pipeline
#             self.text_generator = pipeline(
#                 "text-generation",
#                 model=self.language_model,
#                 tokenizer=self.tokenizer,
#                 device_map='auto' if self.device == 'cuda' else None,
#                 torch_dtype=torch.float16 if self.device == 'cuda' else torch.float32,
#                 pad_token_id=self.tokenizer.eos_token_id
#             )
#             print("✓ Language model loaded successfully")
            
#         except Exception as e:
#             print(f"❌ Error loading language model: {e}")
#             print("Falling back to text-only mode (no answer generation)")
#             self.tokenizer = None
#             self.language_model = None
#             self.text_generator = None
        
#         # Load vector database
#         self.vector_db = self.load_vector_db()
        
#         # Results storage
#         self.results = []
#         self.detailed_results = []
        
#     def load_vector_db(self) -> List:
#         """Load the cached vector database or create if it doesn't exist"""
#         if os.path.exists(VECTOR_DB_FILE):
#             try:
#                 with open(VECTOR_DB_FILE, 'rb') as f:
#                     vector_db = pickle.load(f)
#                 print(f"Loaded vector database with {len(vector_db)} chunks")
#                 return vector_db
#             except Exception as e:
#                 print(f"Error loading vector database: {e}")
#                 print("Will create new vector database...")
        
#         print("Vector database not found. Creating new embeddings...")
#         return self.create_vector_db()
    
#     def read_pdf(self, file_path: str) -> str:
#         """Extract text from PDF file"""
#         try:
#             import PyPDF2
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 text = ""
#                 for page in pdf_reader.pages:
#                     try:
#                         page_text = page.extract_text()
#                         if page_text:
#                             text += page_text + "\n"
#                     except Exception as e:
#                         print(f"Error reading page from {file_path}: {e}")
#                         continue
#             return text
#         except Exception as e:
#             print(f"Error reading PDF {file_path}: {e}")
#             return ""
    
#     def read_docx(self, file_path: str) -> str:
#         """Extract text from DOCX file"""
#         try:
#             from docx import Document
#             doc = Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             print(f"Error reading DOCX {file_path}: {e}")
#             return ""
    
#     def read_txt(self, file_path: str) -> str:
#         """Read text file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 return f.read()
#         except Exception as e:
#             print(f"Error reading TXT {file_path}: {e}")
#             return ""
    
#     def read_json(self, file_path: str) -> str:
#         """Extract text from JSON file with intelligent text extraction"""
#         def extract_text_from_json(obj, path=""):
#             """Recursively extract text from JSON object"""
#             text_parts = []
            
#             if isinstance(obj, dict):
#                 for key, value in obj.items():
#                     current_path = f"{path}.{key}" if path else key
                    
#                     # Extract key names as context
#                     if isinstance(value, (str, int, float)) and value:
#                         text_parts.append(f"{key}: {value}")
#                     elif isinstance(value, (list, dict)):
#                         text_parts.extend(extract_text_from_json(value, current_path))
                        
#             elif isinstance(obj, list):
#                 for i, item in enumerate(obj):
#                     current_path = f"{path}[{i}]" if path else f"[{i}]"
#                     text_parts.extend(extract_text_from_json(item, current_path))
                    
#             elif isinstance(obj, (str, int, float)) and obj:
#                 text_parts.append(str(obj))
                
#             return text_parts
        
#         try:
#             encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
#             for encoding in encodings:
#                 try:
#                     with open(file_path, 'r', encoding=encoding) as file:
#                         data = json.load(file)
                        
#                     # Extract all text from the JSON structure
#                     text_parts = extract_text_from_json(data)
                    
#                     # Join with newlines for better readability
#                     return '\n'.join(text_parts)
                    
#                 except UnicodeDecodeError:
#                     continue
#                 except json.JSONDecodeError as e:
#                     print(f"Invalid JSON format in {file_path}: {e}")
#                     return ""
            
#             print(f"Could not decode JSON file {file_path} with any encoding")
#             return ""
            
#         except Exception as e:
#             print(f"Error reading JSON {file_path}: {e}")
#             return ""
    
#     def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
#         """Split text into overlapping chunks"""
#         words = text.split()
#         chunks = []
        
#         for i in range(0, len(words), chunk_size - overlap):
#             chunk = " ".join(words[i:i + chunk_size])
#             if chunk.strip():
#                 chunks.append(chunk)
        
#         return chunks
    
#     def load_documents(self) -> List[Tuple[str, str, dict]]:
#         """Load all documents from the sources folder"""
#         documents = []
        
#         if not os.path.exists(SOURCES_FOLDER):
#             print(f"Sources folder {SOURCES_FOLDER} not found!")
#             return documents
        
#         for filename in os.listdir(SOURCES_FOLDER):
#             if filename.startswith('~$'):  # Skip temporary files
#                 continue
                
#             file_path = os.path.join(SOURCES_FOLDER, filename)
#             if not os.path.isfile(file_path):
#                 continue
            
#             file_ext = filename.lower().split('.')[-1]
#             text = ""
            
#             print(f"Loading {filename}...")
            
#             if file_ext == 'pdf':
#                 text = self.read_pdf(file_path)
#             elif file_ext == 'docx':
#                 text = self.read_docx(file_path)
#             elif file_ext in ['txt', 'md']:
#                 text = self.read_txt(file_path)
#             elif file_ext == 'json':
#                 text = self.read_json(file_path)
#             else:
#                 print(f"Unsupported file type: {filename}")
#                 continue
            
#             if text.strip():
#                 # Split into chunks
#                 chunks = self.chunk_text(text)
#                 for i, chunk in enumerate(chunks):
#                     metadata = {
#                         'filename': filename,
#                         'chunk_id': i,
#                         'file_type': file_ext
#                     }
#                     documents.append((chunk, filename, metadata))
#                 print(f"  Created {len(chunks)} chunks from {filename}")
#             else:
#                 print(f"  No text extracted from {filename}")
        
#         return documents
    
#     def create_vector_db(self) -> List:
#         """Create vector database with embeddings"""
#         print("Creating vector database with BAAI/bge-large-en-v1.5...")
        
#         # Load documents
#         documents = self.load_documents()
#         if not documents:
#             print("No documents found to create embeddings!")
#             return []
        
#         print(f"Creating embeddings for {len(documents)} chunks...")
        
#         # Create embeddings in batches
#         vector_db = []
#         batch_size = 32
        
#         for i in tqdm(range(0, len(documents), batch_size), desc="Creating embeddings"):
#             batch = documents[i:i + batch_size]
#             texts = [doc[0] for doc in batch]
            
#             try:
#                 embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                
#                 for j, (text, filename, metadata) in enumerate(batch):
#                     vector_db.append((text, embeddings[j], filename, metadata))
                    
#             except Exception as e:
#                 print(f"Error creating embeddings for batch {i//batch_size}: {e}")
#                 continue
        
#         # Save to cache
#         try:
#             with open(VECTOR_DB_FILE, 'wb') as f:
#                 pickle.dump(vector_db, f)
#             print(f"Saved vector database with {len(vector_db)} chunks to {VECTOR_DB_FILE}")
#         except Exception as e:
#             print(f"Error saving vector database: {e}")
        
#         return vector_db
    
#     def cosine_similarity_batch(self, query_embedding: np.ndarray, embeddings: List[np.ndarray]) -> List[float]:
#         """Calculate cosine similarity between query and all embeddings"""
#         embeddings_matrix = np.array(embeddings)
        
#         # Normalize vectors
#         query_norm = query_embedding / np.linalg.norm(query_embedding)
#         embeddings_norm = embeddings_matrix / np.linalg.norm(embeddings_matrix, axis=1, keepdims=True)
        
#         # Calculate similarities
#         similarities = np.dot(embeddings_norm, query_norm)
#         return similarities.tolist()
    
#     def retrieve(self, query: str, top_n: int = RETRIEVAL_TOP_K) -> List[Tuple[str, float, str, dict]]:
#         """Retrieve most relevant chunks - adapted from your P-RAG.py"""
#         try:
#             query_embedding = self.embedding_model.encode([query])[0]
#         except Exception as e:
#             print(f"Error creating query embedding: {e}")
#             return []
        
#         if not self.vector_db:
#             return []
        
#         # Extract embeddings and metadata
#         embeddings = [item[1] for item in self.vector_db]
        
#         # Calculate similarities
#         similarities = self.cosine_similarity_batch(query_embedding, embeddings)
        
#         # Create results with metadata
#         results = []
#         for i, similarity in enumerate(similarities):
#             chunk_text, _, filename, metadata = self.vector_db[i]
#             results.append((chunk_text, float(similarity), filename, metadata))
        
#         # Sort by similarity in descending order
#         results.sort(key=lambda x: x[1], reverse=True)
        
#         return results[:top_n]
    
#     def generate_answer(self, query: str, bloom_level: str) -> Dict[str, Any]:
#         """Generate answer using RAG - adapted from your P-RAG.py"""
#         # Retrieve relevant context
#         start_time = time.time()
#         retrieved_knowledge = self.retrieve(query, top_n=RETRIEVAL_TOP_K)
#         retrieval_time = time.time() - start_time
        
#         # Prepare context
#         if not retrieved_knowledge:
#             context_chunks = []
#             context_notice = "No supporting documents were retrieved. Please answer based on your own general understanding."
#         else:
#             context_chunks = []
#             for chunk, similarity, filename, metadata in retrieved_knowledge:
#                 context_chunks.append(f"[Source: {filename}] {chunk}")
#             context_notice = ""
        
#         # Get Bloom taxonomy instruction
#         bloom_instruction_data = get_bloom_instruction(bloom_level)
#         if isinstance(bloom_instruction_data, dict):
#             thinking_instruction = bloom_instruction_data.get('instruction', '')
#         else:
#             thinking_instruction = str(bloom_instruction_data) if bloom_instruction_data else ""
        
#         # Create prompt
#         instruction_prompt = f'''You are a highly intelligent assistant with expertise in Palestine-related topics.
#                                 {thinking_instruction}
#                                 {context_notice}
#                                 Do not cite or reference specific documents or filenames.
#                                 Answer in your own words, using the context below only as guidance.
#                                 Keep your answer concise and focused.

#                                 Context:
#                                 {chr(10).join(context_chunks)}
#                                 '''
        
#         # Generate response
#         start_time = time.time()
#         try:
#             if self.text_generator is None:
#                 generated_answer = "Error: Language model not loaded"
#                 generation_success = False
#             else:
#                 # Create the full prompt
#                 full_prompt = f"{instruction_prompt}\n\nUser: {query}\nAssistant:"
                
#                 # Generate response using Hugging Face pipeline
#                 response = self.text_generator(
#                     full_prompt,
#                     max_new_tokens=512,
#                     temperature=0.7,
#                     do_sample=True,
#                     pad_token_id=self.tokenizer.eos_token_id,
#                     truncation=True,
#                     return_full_text=False
#                 )
                
#                 generated_answer = response[0]['generated_text'].strip()
#                 generation_success = True
                
#         except Exception as e:
#             generated_answer = f"Error generating response: {e}"
#             generation_success = False
            
#         generation_time = time.time() - start_time
        
#         return {
#             'generated_answer': generated_answer,
#             'retrieved_chunks': len(retrieved_knowledge),
#             'retrieval_time': retrieval_time,
#             'generation_time': generation_time,
#             'generation_success': generation_success,
#             'context_sources': [filename for _, _, filename, _ in retrieved_knowledge],
#             'similarity_scores': [similarity for _, similarity, _, _ in retrieved_knowledge]
#         }
    
#     def calculate_answer_similarity(self, generated_answer: str, correct_choice: str, all_choices: List[str]) -> Dict[str, Any]:
#         """Calculate similarity between generated answer and choices"""
#         # Simple keyword-based matching
#         generated_lower = generated_answer.lower()
        
#         # Calculate similarity with correct answer
#         correct_lower = correct_choice.lower()
#         correct_words = set(correct_lower.split())
#         generated_words = set(generated_lower.split())
        
#         if len(correct_words) > 0:
#             correct_similarity = len(correct_words.intersection(generated_words)) / len(correct_words)
#         else:
#             correct_similarity = 0.0
        
#         # Calculate similarity with all choices
#         choice_similarities = []
#         for choice in all_choices:
#             choice_lower = choice.lower()
#             choice_words = set(choice_lower.split())
#             if len(choice_words) > 0:
#                 similarity = len(choice_words.intersection(generated_words)) / len(choice_words)
#             else:
#                 similarity = 0.0
#             choice_similarities.append(similarity)
        
#         # Determine if the generated answer is most similar to the correct choice
#         max_similarity_idx = choice_similarities.index(max(choice_similarities))
#         is_correct = max_similarity_idx == 0  # Gold answer is always at index 0
        
#         return {
#             'correct_similarity': correct_similarity,
#             'choice_similarities': choice_similarities,
#             'predicted_choice_idx': max_similarity_idx,
#             'is_correct': is_correct,
#             'confidence': max(choice_similarities)
#         }
    
#     def evaluate_question(self, question_data: Dict[str, Any]) -> Dict[str, Any]:
#         """Evaluate a single question"""
#         question_id = question_data['id']
#         query = question_data['query']
#         choices = question_data['choices']
#         gold_idx = question_data['gold']
#         bloom_level = question_data['Bloom Taxonomy Level']
        
#         print(f"Evaluating question {question_id}: {query[:80]}...")
        
#         # Generate answer using RAG
#         rag_result = self.generate_answer(query, bloom_level)
        
#         # Calculate similarity with choices
#         correct_choice = choices[gold_idx]
#         similarity_result = self.calculate_answer_similarity(
#             rag_result['generated_answer'], 
#             correct_choice, 
#             choices
#         )
        
#         # Combine results
#         result = {
#             'question_id': question_id,
#             'query': query,
#             'choices': choices,
#             'gold_idx': gold_idx,
#             'correct_choice': correct_choice,
#             'bloom_level': bloom_level,
#             'generated_answer': rag_result['generated_answer'],
#             'is_correct': similarity_result['is_correct'],
#             'predicted_choice_idx': similarity_result['predicted_choice_idx'],
#             'confidence': similarity_result['confidence'],
#             'correct_similarity': similarity_result['correct_similarity'],
#             'choice_similarities': similarity_result['choice_similarities'],
#             'retrieved_chunks': rag_result['retrieved_chunks'],
#             'retrieval_time': rag_result['retrieval_time'],
#             'generation_time': rag_result['generation_time'],
#             'generation_success': rag_result['generation_success'],
#             'context_sources': rag_result['context_sources'],
#             'similarity_scores': rag_result['similarity_scores']
#         }
        
#         return result
    
#     def load_benchmark(self, filepath: str) -> List[Dict[str, Any]]:
#         """Load benchmark dataset from JSONL file"""
#         questions = []
#         try:
#             with jsonlines.open(filepath) as reader:
#                 for obj in reader:
#                     questions.append(obj)
#             print(f"Loaded {len(questions)} questions from benchmark")
#             return questions
#         except Exception as e:
#             print(f"Error loading benchmark: {e}")
#             return []
    
#     def run_benchmark(self, benchmark_file: str, output_dir: str, start_idx: int = 0, end_idx: int = None):
#         """Run the complete benchmark evaluation"""
#         print("Starting Palestine RAG Benchmark Evaluation")
#         print("=" * 60)
        
#         # Create output directory
#         os.makedirs(output_dir, exist_ok=True)
        
#         # Load benchmark questions
#         questions = self.load_benchmark(benchmark_file)
#         if not questions:
#             print("No questions loaded. Exiting.")
#             return
        
#         # Handle slice indices
#         if end_idx is None:
#             end_idx = len(questions)
#         questions_slice = questions[start_idx:end_idx]
        
#         print(f"Evaluating questions {start_idx} to {end_idx-1} ({len(questions_slice)} questions)")
#         print(f"Vector database has {len(self.vector_db)} chunks")
#         print(f"Using model: {LANGUAGE_MODEL}")
#         print("=" * 60)
        
#         # Evaluate each question
#         start_time = time.time()
        
#         for i, question_data in enumerate(tqdm(questions_slice, desc="Evaluating questions")):
#             try:
#                 result = self.evaluate_question(question_data)
#                 self.results.append(result)
                
#                 # Save intermediate results every 10 questions
#                 if (i + 1) % 10 == 0:
#                     self.save_intermediate_results(output_dir, start_idx + i + 1)
                    
#             except Exception as e:
#                 print(f"Error evaluating question {question_data.get('id', 'unknown')}: {e}")
#                 continue
        
#         total_time = time.time() - start_time
        
#         # Generate comprehensive analysis
#         self.analyze_results(total_time)
        
#         # Save final results
#         self.save_results(output_dir)
        
#         print(f"\nBenchmark completed in {total_time:.2f} seconds")
#         print(f"Results saved to {output_dir}/")
    
#     def analyze_results(self, total_time: float):
#         """Analyze benchmark results and generate statistics"""
#         if not self.results:
#             return
        
#         total_questions = len(self.results)
#         correct_answers = sum(1 for r in self.results if r['is_correct'])
#         overall_accuracy = correct_answers / total_questions
        
#         # Analyze by Bloom taxonomy level
#         bloom_stats = defaultdict(lambda: {'total': 0, 'correct': 0, 'times': []})
        
#         for result in self.results:
#             bloom_level = result['bloom_level']
#             bloom_stats[bloom_level]['total'] += 1
#             if result['is_correct']:
#                 bloom_stats[bloom_level]['correct'] += 1
#             bloom_stats[bloom_level]['times'].append(
#                 result['retrieval_time'] + result['generation_time']
#             )
        
#         # Calculate per-level accuracy
#         for level in bloom_stats:
#             if bloom_stats[level]['total'] > 0:
#                 bloom_stats[level]['accuracy'] = bloom_stats[level]['correct'] / bloom_stats[level]['total']
#                 bloom_stats[level]['avg_time'] = np.mean(bloom_stats[level]['times'])
#             else:
#                 bloom_stats[level]['accuracy'] = 0.0
#                 bloom_stats[level]['avg_time'] = 0.0
        
#         # Overall statistics
#         retrieval_times = [r['retrieval_time'] for r in self.results]
#         generation_times = [r['generation_time'] for r in self.results]
#         confidences = [r['confidence'] for r in self.results]
        
#         self.analysis = {
#             'overall_accuracy': overall_accuracy,
#             'total_questions': total_questions,
#             'correct_answers': correct_answers,
#             'bloom_level_stats': dict(bloom_stats),
#             'timing_stats': {
#                 'total_time': total_time,
#                 'avg_retrieval_time': np.mean(retrieval_times),
#                 'avg_generation_time': np.mean(generation_times),
#                 'avg_total_time_per_question': np.mean([r + g for r, g in zip(retrieval_times, generation_times)])
#             },
#             'confidence_stats': {
#                 'mean_confidence': np.mean(confidences),
#                 'median_confidence': np.median(confidences),
#                 'min_confidence': np.min(confidences),
#                 'max_confidence': np.max(confidences)
#             },
#             'generation_success_rate': sum(1 for r in self.results if r['generation_success']) / total_questions
#         }
    
#     def save_intermediate_results(self, output_dir: str, current_idx: int):
#         """Save intermediate results during evaluation"""
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         filename = f"intermediate_results_{current_idx}_{timestamp}.json"
#         filepath = os.path.join(output_dir, filename)
        
#         with open(filepath, 'w', encoding='utf-8') as f:
#             json.dump(self.results, f, ensure_ascii=False, indent=2)
    
#     def save_results(self, output_dir: str):
#         """Save comprehensive results in JSON and text formats"""
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
#         # Save detailed results as JSON
#         json_file = os.path.join(output_dir, f"benchmark_results_{timestamp}.json")
#         results_data = {
#             'metadata': {
#                 'timestamp': timestamp,
#                 'model': LANGUAGE_MODEL,
#                 'embedding_model': EMBEDDING_MODEL_NAME,
#                 'benchmark_file': BENCHMARK_FILE,
#                 'total_questions': len(self.results),
#                 'vector_db_size': len(self.vector_db)
#             },
#             'analysis': self.analysis,
#             'detailed_results': self.results
#         }
        
#         with open(json_file, 'w', encoding='utf-8') as f:
#             json.dump(results_data, f, ensure_ascii=False, indent=2)
        
#         # Save summary as text
#         txt_file = os.path.join(output_dir, f"benchmark_summary_{timestamp}.txt")
#         self.save_text_summary(txt_file)
        
#         # Save CSV for easy analysis
#         csv_file = os.path.join(output_dir, f"benchmark_results_{timestamp}.csv")
#         self.save_csv_results(csv_file)
        
#         print(f"Results saved:")
#         print(f"  JSON: {json_file}")
#         print(f"  Summary: {txt_file}")
#         print(f"  CSV: {csv_file}")
    
#     def save_text_summary(self, filepath: str):
#         """Save human-readable summary report"""
#         with open(filepath, 'w', encoding='utf-8') as f:
#             f.write("Palestine RAG Benchmark Evaluation Report\n")
#             f.write("=" * 50 + "\n\n")
            
#             # Configuration
#             f.write("Configuration:\n")
#             f.write(f"  Model: {LANGUAGE_MODEL}\n")
#             f.write(f"  Embedding Model: {EMBEDDING_MODEL_NAME}\n")
#             f.write(f"  Vector DB Size: {len(self.vector_db)} chunks\n")
#             f.write(f"  Retrieval Top-K: {RETRIEVAL_TOP_K}\n\n")
            
#             # Overall Results
#             f.write("Overall Results:\n")
#             f.write(f"  Total Questions: {self.analysis['total_questions']}\n")
#             f.write(f"  Correct Answers: {self.analysis['correct_answers']}\n")
#             f.write(f"  Overall Accuracy: {self.analysis['overall_accuracy']:.2%}\n")
#             f.write(f"  Generation Success Rate: {self.analysis['generation_success_rate']:.2%}\n\n")
            
#             # Timing Statistics
#             f.write("Timing Statistics:\n")
#             timing = self.analysis['timing_stats']
#             f.write(f"  Total Evaluation Time: {timing['total_time']:.2f} seconds\n")
#             f.write(f"  Average Retrieval Time: {timing['avg_retrieval_time']:.3f} seconds\n")
#             f.write(f"  Average Generation Time: {timing['avg_generation_time']:.3f} seconds\n")
#             f.write(f"  Average Total Time per Question: {timing['avg_total_time_per_question']:.3f} seconds\n\n")
            
#             # Confidence Statistics
#             f.write("Confidence Statistics:\n")
#             conf = self.analysis['confidence_stats']
#             f.write(f"  Mean Confidence: {conf['mean_confidence']:.3f}\n")
#             f.write(f"  Median Confidence: {conf['median_confidence']:.3f}\n")
#             f.write(f"  Min Confidence: {conf['min_confidence']:.3f}\n")
#             f.write(f"  Max Confidence: {conf['max_confidence']:.3f}\n\n")
            
#             # Bloom Taxonomy Level Analysis
#             f.write("Performance by Bloom Taxonomy Level:\n")
#             for level, stats in self.analysis['bloom_level_stats'].items():
#                 f.write(f"  {level}:\n")
#                 f.write(f"    Questions: {stats['total']}\n")
#                 f.write(f"    Correct: {stats['correct']}\n")
#                 f.write(f"    Accuracy: {stats['accuracy']:.2%}\n")
#                 f.write(f"    Avg Time: {stats['avg_time']:.3f} seconds\n\n")
            
#             # Sample Results
#             f.write("Sample Results (First 5 Questions):\n")
#             f.write("-" * 50 + "\n")
#             for i, result in enumerate(self.results[:5]):
#                 f.write(f"Question {result['question_id']}:\n")
#                 f.write(f"  Query: {result['query']}\n")
#                 f.write(f"  Correct Choice: {result['correct_choice']}\n")
#                 f.write(f"  Generated Answer: {result['generated_answer'][:200]}...\n")
#                 f.write(f"  Is Correct: {result['is_correct']}\n")
#                 f.write(f"  Confidence: {result['confidence']:.3f}\n")
#                 f.write(f"  Bloom Level: {result['bloom_level']}\n\n")
    
#     def save_csv_results(self, filepath: str):
#         """Save results in CSV format for easy analysis"""
#         import csv
        
#         with open(filepath, 'w', newline='', encoding='utf-8') as f:
#             writer = csv.writer(f)
            
#             # Header
#             writer.writerow([
#                 'question_id', 'bloom_level', 'is_correct', 'confidence',
#                 'retrieval_time', 'generation_time', 'retrieved_chunks',
#                 'generation_success', 'query_length', 'answer_length'
#             ])
            
#             # Data
#             for result in self.results:
#                 writer.writerow([
#                     result['question_id'],
#                     result['bloom_level'],
#                     result['is_correct'],
#                     result['confidence'],
#                     result['retrieval_time'],
#                     result['generation_time'],
#                     result['retrieved_chunks'],
#                     result['generation_success'],
#                     len(result['query']),
#                     len(result['generated_answer'])
#                 ])


# def main():
#     """Main function to run the benchmark"""
#     print("Palestine RAG Benchmarking System")
#     print("=" * 50)
    
#     # Check if benchmark file exists
#     if not os.path.exists(BENCHMARK_FILE):
#         print(f"Benchmark file {BENCHMARK_FILE} not found!")
#         return
    
#     # Initialize benchmarker (will create vector DB if needed)
#     print("Initializing RAG Benchmarker...")
#     benchmarker = RAGBenchmarker()
    
#     if not benchmarker.vector_db:
#         print("Failed to create or load vector database!")
#         return
    
#     if benchmarker.text_generator is None:
#         print("Failed to load language model!")
#         return
    
#     # Ask user for configuration
#     print(f"Benchmark file: {BENCHMARK_FILE}")
#     print(f"Vector database: {VECTOR_DB_FILE}")
#     print(f"Output directory: {OUTPUT_DIR}")
#     print(f"Language model: {LANGUAGE_MODEL}")
#     print(f"Embedding model: {EMBEDDING_MODEL_NAME}")
    
#     # Allow user to specify range
#     start_idx = 0
#     end_idx = None
    
#     try:
#         user_start = input(f"Start from question index (default 0): ").strip()
#         if user_start:
#             start_idx = int(user_start)
        
#         user_end = input(f"End at question index (default: all): ").strip()
#         if user_end:
#             end_idx = int(user_end)
#     except ValueError:
#         print("Invalid input, using defaults")
    
#     # Run benchmark
#     try:
#         benchmarker.run_benchmark(BENCHMARK_FILE, OUTPUT_DIR, start_idx, end_idx)
#     except KeyboardInterrupt:
#         print("\nBenchmark interrupted by user")
#         # Save partial results
#         if benchmarker.results:
#             benchmarker.save_results(OUTPUT_DIR)
#     except Exception as e:
#         print(f"Error during benchmark: {e}")
#         # Save partial results
#         if benchmarker.results:
#             benchmarker.save_results(OUTPUT_DIR)


# if __name__ == "__main__":
#     main()











#!/usr/bin/env python3
"""
Palestine RAG Benchmarking Script

This script evaluates the Palestine RAG system against a benchmark dataset containing
questions with multiple choice answers and Bloom taxonomy classifications.

Features:
- Loads benchmark from JSONL file
- Uses existing RAG retrieval and generation
- Evaluates accuracy by comparing generated answers to choices
- Provides detailed analysis by Bloom taxonomy level
- Outputs comprehensive JSON and text reports
"""

import json
import jsonlines
import os
import sys
import time
from datetime import datetime
from typing import List, Dict, Any, Tuple
import re
from collections import defaultdict, Counter
import torch
from sentence_transformers import SentenceTransformer
import numpy as np
import pickle
from tqdm import tqdm
import PyPDF2
from docx import Document as DocxDocument
import csv

# Hugging Face imports
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline



# Import from your existing P-RAG system
from bloom_utils import bloom_levels, get_bloom_instruction

# Configuration - matching your P-RAG.py
SOURCES_FOLDER = "./Sources"
LANGUAGE_MODEL = 'Qwen/Qwen3-4B'  
EMBEDDING_MODEL_NAME = 'intfloat/e5-large-v2'  # Better embedding model
VECTOR_DB_FILE = "vector_db_cache_e5.pkl"  # Separate cache for BGE model

# Benchmark configuration
BENCHMARK_FILE = "converted_benchmark_with_bloom.jsonl"
OUTPUT_DIR = "benchmark_results"
MAX_RESPONSE_LENGTH = 1000
RETRIEVAL_TOP_K = 5

class RAGBenchmarker:
    def __init__(self):
        """Initialize the RAG benchmarker with existing embeddings and models"""
        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        print(f"Using device: {self.device}")
        
        # Initialize embedding model
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME, device=self.device)
        
        # Initialize Hugging Face language model
        print(f"Loading language model: {LANGUAGE_MODEL}")
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(LANGUAGE_MODEL, trust_remote_code=True)
            self.language_model = AutoModelForCausalLM.from_pretrained(
                LANGUAGE_MODEL,
                torch_dtype=torch.float16 if self.device == 'cuda' else torch.float32,
                device_map='auto' if self.device == 'cuda' else None,
                trust_remote_code=True
            )
            
            # Set padding token if not exists
            if self.tokenizer.pad_token is None:
                self.tokenizer.pad_token = self.tokenizer.eos_token
            
            # Create text generation pipeline
            self.text_generator = pipeline(
                "text-generation",
                model=self.language_model,
                tokenizer=self.tokenizer,
                device_map='auto' if self.device == 'cuda' else None,
                torch_dtype=torch.float16 if self.device == 'cuda' else torch.float32,
                pad_token_id=self.tokenizer.eos_token_id
            )
            print("✓ Language model loaded successfully")
            
        except Exception as e:
            print(f"❌ Error loading language model: {e}")
            print("Falling back to text-only mode (no answer generation)")
            self.tokenizer = None
            self.language_model = None
            self.text_generator = None
        
        # Load vector database
        self.vector_db = self.load_vector_db()
        
        # Results storage
        self.results = []
        self.detailed_results = []
        
    def load_vector_db(self) -> List:
        """Load the cached vector database or create if it doesn't exist"""
        if os.path.exists(VECTOR_DB_FILE):
            try:
                with open(VECTOR_DB_FILE, 'rb') as f:
                    vector_db = pickle.load(f)
                print(f"Loaded vector database with {len(vector_db)} chunks")
                return vector_db
            except Exception as e:
                print(f"Error loading vector database: {e}")
                print("Will create new vector database...")
        
        print("Vector database not found. Creating new embeddings...")
        return self.create_vector_db()
    
    def read_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Error reading page from {file_path}: {e}")
                        continue
            return text
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def read_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def read_txt(self, file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def read_json(self, file_path: str) -> str:
        """Extract text from JSON file with intelligent text extraction"""
        def extract_text_from_json(obj, path=""):
            """Recursively extract text from JSON object"""
            text_parts = []
            
            if isinstance(obj, dict):
                for key, value in obj.items():
                    current_path = f"{path}.{key}" if path else key
                    
                    # Extract key names as context
                    if isinstance(value, (str, int, float)) and value:
                        text_parts.append(f"{key}: {value}")
                    elif isinstance(value, (list, dict)):
                        text_parts.extend(extract_text_from_json(value, current_path))
                        
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    current_path = f"{path}[{i}]" if path else f"[{i}]"
                    text_parts.extend(extract_text_from_json(item, current_path))
                    
            elif isinstance(obj, (str, int, float)) and obj:
                text_parts.append(str(obj))
                
            return text_parts
        
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        data = json.load(file)
                        
                    # Extract all text from the JSON structure
                    text_parts = extract_text_from_json(data)
                    
                    # Join with newlines for better readability
                    return '\n'.join(text_parts)
                    
                except UnicodeDecodeError:
                    continue
                except json.JSONDecodeError as e:
                    print(f"Invalid JSON format in {file_path}: {e}")
                    return ""
            
            print(f"Could not decode JSON file {file_path} with any encoding")
            return ""
            
        except Exception as e:
            print(f"Error reading JSON {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks
    
    def load_documents(self) -> List[Tuple[str, str, dict]]:
        """Load all documents from the sources folder"""
        documents = []
        
        if not os.path.exists(SOURCES_FOLDER):
            print(f"Sources folder {SOURCES_FOLDER} not found!")
            return documents
        
        for filename in os.listdir(SOURCES_FOLDER):
            if filename.startswith('~$'):  # Skip temporary files
                continue
                
            file_path = os.path.join(SOURCES_FOLDER, filename)
            if not os.path.isfile(file_path):
                continue
            
            file_ext = filename.lower().split('.')[-1]
            text = ""
            
            print(f"Loading {filename}...")
            
            if file_ext == 'pdf':
                text = self.read_pdf(file_path)
            elif file_ext == 'docx':
                text = self.read_docx(file_path)
            elif file_ext in ['txt', 'md']:
                text = self.read_txt(file_path)
            elif file_ext == 'json':
                text = self.read_json(file_path)
            else:
                print(f"Unsupported file type: {filename}")
                continue
            
            if text.strip():
                # Split into chunks
                chunks = self.chunk_text(text)
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'filename': filename,
                        'chunk_id': i,
                        'file_type': file_ext
                    }
                    documents.append((chunk, filename, metadata))
                print(f"  Created {len(chunks)} chunks from {filename}")
            else:
                print(f"  No text extracted from {filename}")
        
        return documents
    
    def create_vector_db(self) -> List:
        """Create vector database with embeddings"""
        print("Creating vector database with BAAI/bge-large-en-v1.5...")
        
        # Load documents
        documents = self.load_documents()
        if not documents:
            print("No documents found to create embeddings!")
            return []
        
        print(f"Creating embeddings for {len(documents)} chunks...")
        
        # Create embeddings in batches
        vector_db = []
        batch_size = 32
        
        for i in tqdm(range(0, len(documents), batch_size), desc="Creating embeddings"):
            batch = documents[i:i + batch_size]
            texts = [doc[0] for doc in batch]
            
            try:
                embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                
                for j, (text, filename, metadata) in enumerate(batch):
                    vector_db.append((text, embeddings[j], filename, metadata))
                    
            except Exception as e:
                print(f"Error creating embeddings for batch {i//batch_size}: {e}")
                continue
        
        # Save to cache
        try:
            with open(VECTOR_DB_FILE, 'wb') as f:
                pickle.dump(vector_db, f)
            print(f"Saved vector database with {len(vector_db)} chunks to {VECTOR_DB_FILE}")
        except Exception as e:
            print(f"Error saving vector database: {e}")
        
        return vector_db
    
    def cosine_similarity_batch(self, query_embedding: np.ndarray, embeddings: List[np.ndarray]) -> List[float]:
        """Calculate cosine similarity between query and all embeddings"""
        embeddings_matrix = np.array(embeddings)
        
        # Normalize vectors
        query_norm = query_embedding / np.linalg.norm(query_embedding)
        embeddings_norm = embeddings_matrix / np.linalg.norm(embeddings_matrix, axis=1, keepdims=True)
        
        # Calculate similarities
        similarities = np.dot(embeddings_norm, query_norm)
        return similarities.tolist()
    
    def retrieve(self, query: str, top_n: int = RETRIEVAL_TOP_K) -> List[Tuple[str, float, str, dict]]:
        """Retrieve most relevant chunks - adapted from your P-RAG.py"""
        try:
            query_embedding = self.embedding_model.encode([query])[0]
        except Exception as e:
            print(f"Error creating query embedding: {e}")
            return []
        
        if not self.vector_db:
            return []
        
        # Extract embeddings and metadata
        embeddings = [item[1] for item in self.vector_db]
        
        # Calculate similarities
        similarities = self.cosine_similarity_batch(query_embedding, embeddings)
        
        # Create results with metadata
        results = []
        for i, similarity in enumerate(similarities):
            chunk_text, _, filename, metadata = self.vector_db[i]
            results.append((chunk_text, float(similarity), filename, metadata))
        
        # Sort by similarity in descending order
        results.sort(key=lambda x: x[1], reverse=True)
        
        return results[:top_n]
    
    def generate_answer(self, query: str, bloom_level: str, choices: List[str]) -> Dict[str, Any]:
        """Generate answer using RAG with multiple choice selection"""
        # Retrieve relevant context
        start_time = time.time()
        retrieved_knowledge = self.retrieve(query, top_n=RETRIEVAL_TOP_K)
        retrieval_time = time.time() - start_time
        
        # Prepare context
        if not retrieved_knowledge:
            context_chunks = []
            context_notice = "No supporting documents were retrieved. Please answer based on your own general understanding."
        else:
            context_chunks = []
            for chunk, similarity, filename, metadata in retrieved_knowledge:
                context_chunks.append(f"[Source: {filename}] {chunk}")
            context_notice = ""
        
        # Get Bloom taxonomy instruction
        bloom_instruction_data = get_bloom_instruction(bloom_level)
        if isinstance(bloom_instruction_data, dict):
            thinking_instruction = bloom_instruction_data.get('instruction', '')
        else:
            thinking_instruction = str(bloom_instruction_data) if bloom_instruction_data else ""
        
        # Format choices
        formatted_choices = ""
        for i, choice in enumerate(choices):
            formatted_choices += f"{chr(65+i)}. {choice}\n"
        
        # Create multiple choice prompt
        instruction_prompt = f'''You are a highly intelligent assistant with expertise in Palestine-related topics.
{thinking_instruction}
{context_notice}

Based on the context provided below, select the BEST answer from the multiple choice options.
You must respond with ONLY the letter (A, B, C, or D) of the correct answer.
Do not provide any explanation or additional text.

Context:
{chr(10).join(context_chunks)}

Question: {query}

Answer choices:
{formatted_choices}

Your answer (letter only):'''
        
        # Generate response
        start_time = time.time()
        try:
            if self.text_generator is None:
                generated_answer = "Error: Language model not loaded"
                selected_choice_idx = -1
                generation_success = False
            else:
                # Generate response using Hugging Face pipeline
                response = self.text_generator(
                    instruction_prompt,
                    max_new_tokens=10,  # Only need 1-2 tokens for letter response
                    temperature=0.1,   # Lower temperature for more deterministic selection
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id,
                    truncation=True,
                    return_full_text=False
                )
                
                generated_answer = response[0]['generated_text'].strip()
                
                # Extract the selected choice
                selected_choice_idx = self.extract_choice_from_response(generated_answer)
                generation_success = True
                
        except Exception as e:
            generated_answer = f"Error generating response: {e}"
            selected_choice_idx = -1
            generation_success = False
            
        generation_time = time.time() - start_time
        
        return {
            'generated_answer': generated_answer,
            'selected_choice_idx': selected_choice_idx,
            'retrieved_chunks': len(retrieved_knowledge),
            'retrieval_time': retrieval_time,
            'generation_time': generation_time,
            'generation_success': generation_success,
            'context_sources': [filename for _, _, filename, _ in retrieved_knowledge],
            'similarity_scores': [similarity for _, similarity, _, _ in retrieved_knowledge]
        }
    
    def extract_choice_from_response(self, response: str) -> int:
        """Extract choice index from model response"""
        response = response.strip().upper()
        
        # Look for A, B, C, D in the response
        choice_letters = ['A', 'B', 'C', 'D']
        
        # First, try to find a standalone letter
        for i, letter in enumerate(choice_letters):
            if response == letter:
                return i
        
        # Then, look for letter followed by period or space
        for i, letter in enumerate(choice_letters):
            if f"{letter}." in response or f"{letter} " in response:
                return i
        
        # Look for the first occurrence of any choice letter
        for i, letter in enumerate(choice_letters):
            if letter in response:
                return i
        
        # If no valid choice found, return -1 (invalid)
        return -1
    
    def calculate_answer_accuracy(self, selected_choice_idx: int, gold_idx: int, choices: List[str]) -> Dict[str, Any]:
        """Calculate accuracy based on selected choice"""
        # Check if the selection is valid and correct
        is_correct = (selected_choice_idx == gold_idx) and (selected_choice_idx >= 0)
        
        # Calculate confidence (1.0 if valid selection, 0.0 if invalid)
        confidence = 1.0 if selected_choice_idx >= 0 else 0.0
        
        return {
            'is_correct': is_correct,
            'confidence': confidence,
            'predicted_choice_idx': selected_choice_idx,
            'is_valid_selection': selected_choice_idx >= 0
        }
    
    def evaluate_question(self, question_data: Dict[str, Any]) -> Dict[str, Any]:
        """Evaluate a single question"""
        question_id = question_data['id']
        query = question_data['query']
        choices = question_data['choices']
        gold_idx = question_data['gold']
        bloom_level = question_data['Bloom Taxonomy Level']
        
        print(f"Evaluating question {question_id}: {query[:80]}...")
        
        # Generate answer using RAG with multiple choice
        rag_result = self.generate_answer(query, bloom_level, choices)
        
        # Calculate accuracy based on choice selection
        correct_choice = choices[gold_idx]
        accuracy_result = self.calculate_answer_accuracy(
            rag_result['selected_choice_idx'], 
            gold_idx, 
            choices
        )
        
        # Get the selected choice text if valid
        selected_choice_text = ""
        if rag_result['selected_choice_idx'] >= 0 and rag_result['selected_choice_idx'] < len(choices):
            selected_choice_text = choices[rag_result['selected_choice_idx']]
        
        # Combine results
        result = {
            'question_id': question_id,
            'query': query,
            'choices': choices,
            'gold_idx': gold_idx,
            'correct_choice': correct_choice,
            'bloom_level': bloom_level,
            'generated_answer': rag_result['generated_answer'],
            'selected_choice_idx': rag_result['selected_choice_idx'],
            'selected_choice_text': selected_choice_text,
            'is_correct': accuracy_result['is_correct'],
            'confidence': accuracy_result['confidence'],
            'is_valid_selection': accuracy_result['is_valid_selection'],
            'retrieved_chunks': rag_result['retrieved_chunks'],
            'retrieval_time': rag_result['retrieval_time'],
            'generation_time': rag_result['generation_time'],
            'generation_success': rag_result['generation_success'],
            'context_sources': rag_result['context_sources'],
            'similarity_scores': rag_result['similarity_scores']
        }
        
        return result
    
    def load_benchmark(self, filepath: str) -> List[Dict[str, Any]]:
        """Load benchmark dataset from JSONL file"""
        questions = []
        try:
            with jsonlines.open(filepath) as reader:
                for obj in reader:
                    questions.append(obj)
            print(f"Loaded {len(questions)} questions from benchmark")
            return questions
        except Exception as e:
            print(f"Error loading benchmark: {e}")
            return []
    
    def run_benchmark(self, benchmark_file: str, output_dir: str, start_idx: int = 0, end_idx: int = None):
        """Run the complete benchmark evaluation"""
        print("Starting Palestine RAG Benchmark Evaluation")
        print("=" * 60)
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Load benchmark questions
        questions = self.load_benchmark(benchmark_file)
        if not questions:
            print("No questions loaded. Exiting.")
            return
        
        # Handle slice indices
        if end_idx is None:
            end_idx = len(questions)
        questions_slice = questions[start_idx:end_idx]
        
        print(f"Evaluating questions {start_idx} to {end_idx-1} ({len(questions_slice)} questions)")
        print(f"Vector database has {len(self.vector_db)} chunks")
        print(f"Using model: {LANGUAGE_MODEL}")
        print("=" * 60)
        
        # Evaluate each question
        start_time = time.time()
        
        for i, question_data in enumerate(tqdm(questions_slice, desc="Evaluating questions")):
            try:
                result = self.evaluate_question(question_data)
                self.results.append(result)
                
                # Save intermediate results every 10 questions
                if (i + 1) % 10 == 0:
                    self.save_intermediate_results(output_dir, start_idx + i + 1)
                    
            except Exception as e:
                print(f"Error evaluating question {question_data.get('id', 'unknown')}: {e}")
                continue
        
        total_time = time.time() - start_time
        
        # Generate comprehensive analysis
        self.analyze_results(total_time)
        
        # Save final results
        self.save_results(output_dir)
        
        print(f"\nBenchmark completed in {total_time:.2f} seconds")
        print(f"Results saved to {output_dir}/")
    
    def analyze_results(self, total_time: float):
        """Analyze benchmark results and generate statistics"""
        if not self.results:
            return
        
        total_questions = len(self.results)
        correct_answers = sum(1 for r in self.results if r['is_correct'])
        overall_accuracy = correct_answers / total_questions
        
        # Analyze by Bloom taxonomy level
        bloom_stats = defaultdict(lambda: {'total': 0, 'correct': 0, 'times': []})
        
        for result in self.results:
            bloom_level = result['bloom_level']
            bloom_stats[bloom_level]['total'] += 1
            if result['is_correct']:
                bloom_stats[bloom_level]['correct'] += 1
            bloom_stats[bloom_level]['times'].append(
                result['retrieval_time'] + result['generation_time']
            )
        
        # Calculate per-level accuracy
        for level in bloom_stats:
            if bloom_stats[level]['total'] > 0:
                bloom_stats[level]['accuracy'] = bloom_stats[level]['correct'] / bloom_stats[level]['total']
                bloom_stats[level]['avg_time'] = np.mean(bloom_stats[level]['times'])
            else:
                bloom_stats[level]['accuracy'] = 0.0
                bloom_stats[level]['avg_time'] = 0.0
        
        # Overall statistics
        retrieval_times = [r['retrieval_time'] for r in self.results]
        generation_times = [r['generation_time'] for r in self.results]
        confidences = [r['confidence'] for r in self.results]
        valid_selections = sum(1 for r in self.results if r['is_valid_selection'])
        
        self.analysis = {
            'overall_accuracy': overall_accuracy,
            'total_questions': total_questions,
            'correct_answers': correct_answers,
            'valid_selections': valid_selections,
            'valid_selection_rate': valid_selections / total_questions,
            'bloom_level_stats': dict(bloom_stats),
            'timing_stats': {
                'total_time': total_time,
                'avg_retrieval_time': np.mean(retrieval_times),
                'avg_generation_time': np.mean(generation_times),
                'avg_total_time_per_question': np.mean([r + g for r, g in zip(retrieval_times, generation_times)])
            },
            'confidence_stats': {
                'mean_confidence': np.mean(confidences),
                'median_confidence': np.median(confidences),
                'min_confidence': np.min(confidences),
                'max_confidence': np.max(confidences)
            },
            'generation_success_rate': sum(1 for r in self.results if r['generation_success']) / total_questions
        }
    
    def save_intermediate_results(self, output_dir: str, current_idx: int):
        """Save intermediate results during evaluation"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"intermediate_results_{current_idx}_{timestamp}.json"
        filepath = os.path.join(output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.results, f, ensure_ascii=False, indent=2)
    
    def save_results(self, output_dir: str):
        """Save comprehensive results in JSON and text formats"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save detailed results as JSON
        json_file = os.path.join(output_dir, f"benchmark_results_{timestamp}.json")
        results_data = {
            'metadata': {
                'timestamp': timestamp,
                'model': LANGUAGE_MODEL,
                'embedding_model': EMBEDDING_MODEL_NAME,
                'benchmark_file': BENCHMARK_FILE,
                'total_questions': len(self.results),
                'vector_db_size': len(self.vector_db)
            },
            'analysis': self.analysis,
            'detailed_results': self.results
        }
        
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(results_data, f, ensure_ascii=False, indent=2)
        
        # Save summary as text
        txt_file = os.path.join(output_dir, f"benchmark_summary_{timestamp}.txt")
        self.save_text_summary(txt_file)
        
        # Save CSV for easy analysis
        csv_file = os.path.join(output_dir, f"benchmark_results_{timestamp}.csv")
        self.save_csv_results(csv_file)
        
        print(f"Results saved:")
        print(f"  JSON: {json_file}")
        print(f"  Summary: {txt_file}")
        print(f"  CSV: {csv_file}")
    
    def save_text_summary(self, filepath: str):
        """Save human-readable summary report"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("Palestine RAG Benchmark Evaluation Report\n")
            f.write("=" * 50 + "\n\n")
            
            # Configuration
            f.write("Configuration:\n")
            f.write(f"  Model: {LANGUAGE_MODEL}\n")
            f.write(f"  Embedding Model: {EMBEDDING_MODEL_NAME}\n")
            f.write(f"  Vector DB Size: {len(self.vector_db)} chunks\n")
            f.write(f"  Retrieval Top-K: {RETRIEVAL_TOP_K}\n\n")
            
            # Overall Results
            f.write("Overall Results:\n")
            f.write(f"  Total Questions: {self.analysis['total_questions']}\n")
            f.write(f"  Correct Answers: {self.analysis['correct_answers']}\n")
            f.write(f"  Overall Accuracy: {self.analysis['overall_accuracy']:.2%}\n")
            f.write(f"  Valid Selections: {self.analysis['valid_selections']}\n")
            f.write(f"  Valid Selection Rate: {self.analysis['valid_selection_rate']:.2%}\n")
            f.write(f"  Generation Success Rate: {self.analysis['generation_success_rate']:.2%}\n\n")
            
            # Timing Statistics
            f.write("Timing Statistics:\n")
            timing = self.analysis['timing_stats']
            f.write(f"  Total Evaluation Time: {timing['total_time']:.2f} seconds\n")
            f.write(f"  Average Retrieval Time: {timing['avg_retrieval_time']:.3f} seconds\n")
            f.write(f"  Average Generation Time: {timing['avg_generation_time']:.3f} seconds\n")
            f.write(f"  Average Total Time per Question: {timing['avg_total_time_per_question']:.3f} seconds\n\n")
            
            # Confidence Statistics
            f.write("Confidence Statistics:\n")
            conf = self.analysis['confidence_stats']
            f.write(f"  Mean Confidence: {conf['mean_confidence']:.3f}\n")
            f.write(f"  Median Confidence: {conf['median_confidence']:.3f}\n")
            f.write(f"  Min Confidence: {conf['min_confidence']:.3f}\n")
            f.write(f"  Max Confidence: {conf['max_confidence']:.3f}\n\n")
            
            # Bloom Taxonomy Level Analysis
            f.write("Performance by Bloom Taxonomy Level:\n")
            for level, stats in self.analysis['bloom_level_stats'].items():
                f.write(f"  {level}:\n")
                f.write(f"    Questions: {stats['total']}\n")
                f.write(f"    Correct: {stats['correct']}\n")
                f.write(f"    Accuracy: {stats['accuracy']:.2%}\n")
                f.write(f"    Avg Time: {stats['avg_time']:.3f} seconds\n\n")
            
            # Sample Results
            f.write("Sample Results (First 5 Questions):\n")
            f.write("-" * 50 + "\n")
            for i, result in enumerate(self.results[:5]):
                f.write(f"Question {result['question_id']}:\n")
                f.write(f"  Query: {result['query']}\n")
                f.write(f"  Correct Choice: {result['correct_choice']}\n")
                f.write(f"  Selected Choice: {result['selected_choice_text']}\n")
                f.write(f"  Model Response: {result['generated_answer']}\n")
                f.write(f"  Is Correct: {result['is_correct']}\n")
                f.write(f"  Valid Selection: {result['is_valid_selection']}\n")
                f.write(f"  Confidence: {result['confidence']:.3f}\n")
                f.write(f"  Bloom Level: {result['bloom_level']}\n\n")
    
    def save_csv_results(self, filepath: str):
        """Save results in CSV format for easy analysis"""
        import csv
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            # Header
            writer.writerow([
                'question_id', 'bloom_level', 'is_correct', 'confidence',
                'retrieval_time', 'generation_time', 'retrieved_chunks',
                'generation_success', 'query_length', 'answer_length',
                'selected_choice_idx', 'is_valid_selection'
            ])
            
            # Data
            for result in self.results:
                writer.writerow([
                    result['question_id'],
                    result['bloom_level'],
                    result['is_correct'],
                    result['confidence'],
                    result['retrieval_time'],
                    result['generation_time'],
                    result['retrieved_chunks'],
                    result['generation_success'],
                    len(result['query']),
                    len(result['generated_answer']),
                    result['selected_choice_idx'],
                    result['is_valid_selection']
                ])


def main():
    """Main function to run the benchmark"""
    print("Palestine RAG Benchmarking System")
    print("=" * 50)
    
    # Check if benchmark file exists
    if not os.path.exists(BENCHMARK_FILE):
        print(f"Benchmark file {BENCHMARK_FILE} not found!")
        return
    
    # Initialize benchmarker (will create vector DB if needed)
    print("Initializing RAG Benchmarker...")
    benchmarker = RAGBenchmarker()
    
    if not benchmarker.vector_db:
        print("Failed to create or load vector database!")
        return
    
    if benchmarker.text_generator is None:
        print("Failed to load language model!")
        return
    
    # Ask user for configuration
    print(f"Benchmark file: {BENCHMARK_FILE}")
    print(f"Vector database: {VECTOR_DB_FILE}")
    print(f"Output directory: {OUTPUT_DIR}")
    print(f"Language model: {LANGUAGE_MODEL}")
    print(f"Embedding model: {EMBEDDING_MODEL_NAME}")
    
    # Allow user to specify range
    start_idx = 0
    end_idx = None
    
    try:
        user_start = input(f"Start from question index (default 0): ").strip()
        if user_start:
            start_idx = int(user_start)
        
        user_end = input(f"End at question index (default: all): ").strip()
        if user_end:
            end_idx = int(user_end)
    except ValueError:
        print("Invalid input, using defaults")
    
    # Run benchmark
    try:
        benchmarker.run_benchmark(BENCHMARK_FILE, OUTPUT_DIR, start_idx, end_idx)
    except KeyboardInterrupt:
        print("\nBenchmark interrupted by user")
        # Save partial results
        if benchmarker.results:
            benchmarker.save_results(OUTPUT_DIR)
    except Exception as e:
        print(f"Error during benchmark: {e}")
        # Save partial results
        if benchmarker.results:
            benchmarker.save_results(OUTPUT_DIR)


if __name__ == "__main__":
    main()